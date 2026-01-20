import pytesseract
import cv2
from docx import Document
import os
from collections import defaultdict

# Uncomment and set if Tesseract is not in PATH
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

def image_to_word(image_path, output_docx):
    """
    Universal Image → Word converter.
    - Preserves table layout for invoices.
    - Writes paragraphs for letters or simple text.
    """
    # Read image and convert to grayscale
    img = cv2.imread(image_path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    # OCR: get detailed word-level info
    data = pytesseract.image_to_data(gray, output_type=pytesseract.Output.DICT)

    # Group words by approximate y-coordinate → rows
    rows = defaultdict(list)
    for i in range(len(data['text'])):
        text = data['text'][i].strip()
        if not text:
            continue
        x = data['left'][i]
        y = data['top'][i]
        rows[y].append((x, text))

    sorted_rows = sorted(rows.items(), key=lambda x: x[0])
    doc = Document()

    table_created = False
    table = None

    for _, row in sorted_rows:
        # Sort words in the row by x-coordinate (left → right)
        row_cells = [word for x, word in sorted(row, key=lambda x: x[0])]

        # Decide if this is a table row or paragraph
        if len(row_cells) > 2:  # heuristic: 3+ words → table
            if not table_created:
                # Determine max columns for table
                max_cols = max(len([w for _, w in r]) for _, r in sorted_rows)
                table = doc.add_table(rows=0, cols=max_cols)
                table.style = 'Table Grid'
                table_created = True

            # Add new row
            new_row = table.add_row().cells
            for i in range(len(new_row)):
                new_row[i].text = row_cells[i] if i < len(row_cells) else ""
        else:
            # Treat as paragraph if few words
            doc.add_paragraph(" ".join(row_cells))

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_docx), exist_ok=True)
    doc.save(output_docx)
