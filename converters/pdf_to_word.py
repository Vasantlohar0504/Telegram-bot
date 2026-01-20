import pdfplumber
from docx import Document

def pdf_to_word(pdf_path, output_path):
    """
    Convert PDF to Word, preserving text and tables.
    Works even if tables are not perfectly bordered.
    """
    doc = Document()

    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            # 1️⃣ Extract normal text
            text = page.extract_text()
            if text:
                doc.add_paragraph(f"Text from page {page_num}:")
                for line in text.split("\n"):
                    doc.add_paragraph(line)
                doc.add_paragraph("\n")

            # 2️⃣ Extract tables (if any)
            tables = page.extract_tables()
            for table in tables:
                rows, cols = len(table), len(table[0])
                word_table = doc.add_table(rows=rows, cols=cols)
                word_table.style = "Table Grid"

                for r in range(rows):
                    for c in range(cols):
                        word_table.rows[r].cells[c].text = table[r][c] if table[r][c] else ""

                doc.add_paragraph("\n")

    doc.save(output_path)
    print(f"PDF converted to Word: {output_path}")
